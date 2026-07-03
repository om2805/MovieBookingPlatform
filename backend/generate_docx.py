import re
import os
from docx import Document
from docx.shared import Pt

ROOT = r"C:\Users\omupg\IdeaProjects\bms"
SRC = os.path.join(ROOT, 'src', 'main', 'java', 'org', 'example', 'bms')
OUT_DIR = os.path.join(ROOT, 'docs')
os.makedirs(OUT_DIR, exist_ok=True)
OUT_PATH = os.path.join(OUT_DIR, 'backend_documentation.docx')

def read_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception:
        return ''


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def summarize_model(content):
    fields = re.findall(r"private\s+([\w<>\[\]]+)\s+(\w+);", content)
    lines = [f"{t} {n}" for t,n in fields]
    return lines


def extract_mappings(content):
    mappings = []
    lines = content.splitlines()
    current_mapping = None
    for i,l in enumerate(lines):
        m = re.search(r"@(GetMapping|PostMapping|PutMapping|DeleteMapping|RequestMapping)\(([^)]*)\)", l.strip())
        if m:
            annot = m.group(1)
            path = m.group(2)
            # next non-empty line likely method signature
            sig = ''
            for j in range(i+1, min(i+6, len(lines))):
                if lines[j].strip():
                    sig = lines[j].strip()
                    break
            mappings.append((annot, path, sig))
    return mappings


def extract_class_name(content):
    m = re.search(r"public\s+class\s+(\w+)", content)
    return m.group(1) if m else None


doc = Document()
add_heading(doc, 'Backend Documentation - BookMyShow (bms)', level=1)
add_paragraph(doc, 'Generated summary of backend codebase. Contains architecture, endpoints, models, services, repositories, security and runtime configuration.')

# Application & config
app_file = os.path.join(ROOT, 'src', 'main', 'java', 'org', 'example', 'bms', 'BmsApplication.java')
app = read_file(app_file)
add_heading(doc, 'Application entrypoint', level=2)
add_paragraph(doc, f'File: {app_file}')
if app:
    add_paragraph(doc, 'Spring Boot application with class: ' + (extract_class_name(app) or 'BmsApplication'))

# application.properties
props = read_file(os.path.join(ROOT, 'src', 'main', 'resources', 'application.properties'))
add_heading(doc, 'Configuration (application.properties)', level=2)
add_paragraph(doc, props or 'No properties found')

# Security
sec_file = os.path.join(SRC, 'configuration', 'SecurityConfig.java')
sec = read_file(sec_file)
add_heading(doc, 'Security configuration', level=2)
add_paragraph(doc, f'File: {sec_file}')
if sec:
    add_paragraph(doc, 'Highlights:')
    add_paragraph(doc, '- JWT filter added before UsernamePasswordAuthenticationFilter')
    add_paragraph(doc, '- Stateless session management (SessionCreationPolicy.STATELESS)')
    add_paragraph(doc, '- Public endpoints: /api/auth/**, GET /api/movies/**, GET /api/shows/**, GET /api/theaters/**')
    add_paragraph(doc, '- Role-based access rules for USER, MANAGER, ADMIN')

# Jwt Service
jwt_file = os.path.join(SRC, 'security', 'JwtService.java')
jwt = read_file(jwt_file)
add_heading(doc, 'JWT handling', level=2)
add_paragraph(doc, f'File: {jwt_file}')
add_paragraph(doc, '- Uses jjwt to generate and validate tokens; secret and expiration from application.properties (jwt.secret, jwt.expiration)')

# Controllers
add_heading(doc, 'Controllers and Endpoints', level=2)
controllers_dir = os.path.join(SRC, 'controller')
for fname in os.listdir(controllers_dir):
    if not fname.endswith('.java'):
        continue
    path = os.path.join(controllers_dir, fname)
    content = read_file(path)
    class_name = extract_class_name(content) or fname.replace('.java','')
    add_heading(doc, class_name, level=3)
    add_paragraph(doc, f'File: {path}')
    mappings = extract_mappings(content)
    if mappings:
        add_paragraph(doc, 'Endpoints:')
        for annot, p, sig in mappings:
            add_paragraph(doc, f'- {annot} {p} -> {sig}')
    else:
        add_paragraph(doc, 'No mapping annotations found (or could not parse).')

# Services
add_heading(doc, 'Services', level=2)
services_dir = os.path.join(SRC, 'service')
for fname in os.listdir(services_dir):
    if not fname.endswith('.java'):
        continue
    path = os.path.join(services_dir, fname)
    content = read_file(path)
    class_name = extract_class_name(content) or fname.replace('.java','')
    add_heading(doc, class_name, level=3)
    add_paragraph(doc, f'File: {path}')
    # list public methods
    methods = re.findall(r"public\s+[\w<>,\s\[\]]+\s+(\w+)\s*\([^{;]*\)", content)
    if methods:
        add_paragraph(doc, 'Public methods:')
        for m in methods:
            add_paragraph(doc, f'- {m}')

# Repositories
add_heading(doc, 'Repositories', level=2)
repo_dir = os.path.join(SRC, 'repo')
for fname in os.listdir(repo_dir):
    if not fname.endswith('.java'):
        continue
    path = os.path.join(repo_dir, fname)
    content = read_file(path)
    add_heading(doc, fname.replace('.java',''), level=3)
    add_paragraph(doc, f'File: {path}')
    # custom method signatures
    methods = re.findall(r"\n\s*(\w[\w<>_]+\s+\w+\([^)]*\));", content)
    # simpler: list method names present by patterns
    custom = re.findall(r"\n\s*([\w<>\[\]]+\s+\w+\([^)]*\));", content)
    if custom:
        add_paragraph(doc, 'Declared methods:')
        for c in custom:
            add_paragraph(doc, f'- {c.strip()}')

# Models
add_heading(doc, 'Data models (Entities)', level=2)
model_dir = os.path.join(SRC, 'model')
for fname in os.listdir(model_dir):
    if not fname.endswith('.java'):
        continue
    path = os.path.join(model_dir, fname)
    content = read_file(path)
    class_name = extract_class_name(content) or fname.replace('.java','')
    add_heading(doc, class_name, level=3)
    add_paragraph(doc, f'File: {path}')
    fields = summarize_model(content)
    if fields:
        add_paragraph(doc, 'Fields:')
        for f in fields:
            add_paragraph(doc, f'- {f}')

# Concurrency & Transactions
add_heading(doc, 'Concurrency, Locking and Transactions', level=2)
add_paragraph(doc, 'BookingService.createBooking is @Transactional and the ShowSeatRepository.findSeatsForUpdate uses PESSIMISTIC_WRITE lock. This prevents double-booking by acquiring DB-level locks on selected seats during booking flow.')

# Exception handling
add_heading(doc, 'Exception handling', level=2)
add_paragraph(doc, 'GlobalExceptionHandler centralizes handling for ResourceNotFoundException, SeatUnavailableException, BadCredentialsException, AccessDeniedException and generic exceptions, returning ErrorResponse DTO with timestamp, message and HTTP status.')

# How to build and run
add_heading(doc, 'Build and run', level=2)
add_paragraph(doc, '- This is a Maven Spring Boot project (pom.xml present).')
add_paragraph(doc, '- Build: mvn clean package')
add_paragraph(doc, '- Run: java -jar target/*.jar or mvn spring-boot:run')
add_paragraph(doc, '- Requires MySQL configured in application.properties (spring.datasource.*).')

# Notes & Recommendations
add_heading(doc, 'Notes and Recommendations', level=2)
add_paragraph(doc, '- Consider externalizing jwt.secret via environment variable or vault; current secret in properties is insecure for production.')
add_paragraph(doc, '- Consider better error messages and avoiding RuntimeException in AuthService.signup when email exists; use a custom exception.')
add_paragraph(doc, '- Add unit/integration tests for booking concurrency and refund logic.')

# Save
doc.save(OUT_PATH)
print('DOCX generated at', OUT_PATH)
